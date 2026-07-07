import difflib
import io
import PyPDF2
import docx
import openpyxl
from fpdf import FPDF
import uuid
import os

def extract_advanced_text(file, ext="") -> str:
    if ext == '.pdf':
        pdf_reader = PyPDF2.PdfReader(file)
        text = ''
        for page in pdf_reader.pages:
            t = page.extract_text()
            if t:
                text += t + '\n'
        return text
    elif ext == '.docx':
        doc = docx.Document(file)
        return '\n'.join([para.text for para in doc.paragraphs])
    elif ext == '.xlsx':
        wb = openpyxl.load_workbook(file)
        sheet = wb.active
        lines = []
        for row in sheet.iter_rows(values_only=True):
            cleaned = [str(cell) if cell is not None else "" for cell in row]
            lines.append('\t'.join(cleaned))
        return '\n'.join(lines)
    else:
        return file.read().decode('utf-8')

def generate_advanced_structured_diff(text1: str, text2: str):
    lines1 = [x for x in (text1.replace('\r\n', '\n').split('\n')) if x]
    lines2 = [x for x in (text2.replace('\r\n', '\n').split('\n')) if x]

    # Calculate overall similarity
    overall_sm = difflib.SequenceMatcher(None, text1, text2)
    overall_similarity = round(overall_sm.ratio() * 100)

    added = []
    removed = []
    modified = []

    sm = difflib.SequenceMatcher(None, lines1, lines2)
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == 'replace':
            max_len = max(i2 - i1, j2 - j1)
            for k in range(max_len):
                old_line = lines1[i1+k] if (i1+k) < i2 else None
                new_line = lines2[j1+k] if (j1+k) < j2 else None

                if old_line is not None and new_line is not None:
                    line_sm = difflib.SequenceMatcher(None, old_line, new_line)
                    sim_pct = round(line_sm.ratio() * 100)
                    if sim_pct >= 70:
                        modified.append((j1+k+1, old_line, new_line, sim_pct))
                    else:
                        removed.append((i1+k+1, old_line))
                        added.append((j1+k+1, new_line))
                elif old_line is not None:
                    removed.append((i1+k+1, old_line))
                elif new_line is not None:
                    added.append((j1+k+1, new_line))
        elif tag == 'delete':
            for k, line in enumerate(lines1[i1:i2]):
                removed.append((i1+k+1, line))
        elif tag == 'insert':
            for k, line in enumerate(lines2[j1:j2]):
                added.append((j1+k+1, line))

    return {
        'overall_similarity': overall_similarity,
        'added': added,
        'removed': removed,
        'modified': modified,
        'summary': {
            'added': len(added),
            'removed': len(removed),
            'modified': len(modified)
        }
    }

def insert_line_and_generate(file, ext, line_number: int, new_content: str):
    """
    Inserts a new line at `line_number` (1-based) into the document,
    regenerating it into a binary BytesIO object.
    """
    buf = io.BytesIO()
    mime = 'text/plain'
    
    if ext == '.pdf':
        text = extract_advanced_text(file, ext)
        lines = text.split('\n')
        # PDF line insertion
        idx = max(0, line_number - 1)
        if idx >= len(lines):
            lines.append(new_content)
        else:
            lines.insert(idx, new_content)

        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=11)
        # Note: handling fpdf unicode issues properly is hard without advanced fonts.
        # Fallback to ascii/latin1 filtering to prevent crash
        for row in lines:
            safe_text = row.encode('latin-1', 'replace').decode('latin-1')
            pdf.cell(200, 10, txt=safe_text, ln=1)
        
        tmp_name = os.path.join("C:\\Users\\PONNADA HARSHA\\.gemini\\antigravity\\scratch", str(uuid.uuid4())+".pdf")
        pdf.output(tmp_name)
        with open(tmp_name, 'rb') as f:
            buf.write(f.read())
        os.remove(tmp_name)
        buf.seek(0)
        mime = 'application/pdf'
    
    elif ext == '.docx':
        doc = docx.Document(file)
        new_doc = docx.Document()
        paras = [p.text for p in doc.paragraphs]
        idx = max(0, line_number - 1)
        if idx >= len(paras):
            paras.append(new_content)
        else:
            paras.insert(idx, new_content)
            
        for p in paras:
            new_doc.add_paragraph(p)
            
        new_doc.save(buf)
        buf.seek(0)
        mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        
    elif ext == '.xlsx':
        wb = openpyxl.load_workbook(file)
        sheet = wb.active
        idx = max(1, line_number)
        sheet.insert_rows(idx)
        
        # basic tab splits
        cols = new_content.split('\t')
        for i, val in enumerate(cols):
            sheet.cell(row=idx, column=i+1).value = val
            
        wb.save(buf)
        buf.seek(0)
        mime = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        
    else: 
        text = extract_advanced_text(file, '')
        lines = text.split('\n')
        idx = max(0, line_number - 1)
        if idx >= len(lines):
            lines.append(new_content)
        else:
            lines.insert(idx, new_content)
            
        final_text = '\n'.join(lines)
        buf.write(final_text.encode('utf-8'))
        buf.seek(0)
        mime = 'text/plain'
        
    return buf, mime
